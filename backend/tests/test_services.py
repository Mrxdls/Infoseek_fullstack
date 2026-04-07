"""
Tests for service layer — vector store, S3, session, auth, extractor, celery tasks.
All external dependencies are mocked.
"""

import hashlib
import json
import uuid
from datetime import datetime, timedelta, timezone
from unittest.mock import AsyncMock, MagicMock, patch, PropertyMock

import numpy as np
import pytest
import pytest_asyncio
from httpx import AsyncClient
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models.models import (
    Conversation, Document, DocumentChunk, DocumentStatus, DocumentType,
    Message, MessageRole, SessionType, User, UserRole,
)


# ─── Vector Store Service Tests ──────────────────────────────────────────────


class TestVectorStoreService:
    @patch("app.services.ingestion.vector_store.QdrantClient")
    @patch("app.services.ingestion.vector_store.openai_client")
    def test_embed_texts(self, mock_openai, mock_qdrant_cls):
        from app.services.ingestion.vector_store import VectorStoreService

        # Mock collection check
        mock_client = MagicMock()
        mock_client.get_collections.return_value = MagicMock(collections=[MagicMock(name="rag_chunks")])
        mock_qdrant_cls.return_value = mock_client

        # Mock OpenAI embeddings
        mock_embedding = MagicMock()
        mock_embedding.embedding = [0.1] * 1536
        mock_openai.embeddings.create.return_value = MagicMock(data=[mock_embedding])

        vs = VectorStoreService()
        result = vs.embed_texts(["hello"])
        assert len(result) == 1
        assert len(result[0]) == 1536

    @patch("app.services.ingestion.vector_store.QdrantClient")
    @patch("app.services.ingestion.vector_store.openai_client")
    def test_embed_query(self, mock_openai, mock_qdrant_cls):
        from app.services.ingestion.vector_store import VectorStoreService

        mock_client = MagicMock()
        mock_client.get_collections.return_value = MagicMock(collections=[MagicMock(name="rag_chunks")])
        mock_qdrant_cls.return_value = mock_client

        mock_embedding = MagicMock()
        mock_embedding.embedding = [0.2] * 1536
        mock_openai.embeddings.create.return_value = MagicMock(data=[mock_embedding])

        vs = VectorStoreService()
        result = vs.embed_query("test query")
        assert len(result) == 1536

    @patch("app.services.ingestion.vector_store.QdrantClient")
    @patch("app.services.ingestion.vector_store.openai_client")
    def test_ensure_collection_creates_new(self, mock_openai, mock_qdrant_cls):
        from app.services.ingestion.vector_store import VectorStoreService

        mock_client = MagicMock()
        mock_client.get_collections.return_value = MagicMock(collections=[])
        mock_qdrant_cls.return_value = mock_client

        vs = VectorStoreService()
        mock_client.create_collection.assert_called_once()
        assert mock_client.create_payload_index.call_count == 3

    @patch("app.services.ingestion.vector_store.QdrantClient")
    @patch("app.services.ingestion.vector_store.openai_client")
    def test_upsert_chunks(self, mock_openai, mock_qdrant_cls):
        from app.services.ingestion.vector_store import VectorStoreService

        mock_client = MagicMock()
        mock_client.get_collections.return_value = MagicMock(collections=[MagicMock(name="rag_chunks")])
        mock_qdrant_cls.return_value = mock_client

        mock_emb = MagicMock()
        mock_emb.embedding = [0.1] * 1536
        mock_openai.embeddings.create.return_value = MagicMock(data=[mock_emb, mock_emb])

        vs = VectorStoreService()
        chunks = [
            {"chunk_id": "c1", "chunk_text": "text1", "document_id": "d1", "document_type": "learning_material"},
            {"chunk_id": "c2", "chunk_text": "text2", "document_id": "d1", "document_type": "learning_material"},
        ]
        ids = vs.upsert_chunks(chunks)
        assert len(ids) == 2
        mock_client.upsert.assert_called_once()

    @patch("app.services.ingestion.vector_store.QdrantClient")
    @patch("app.services.ingestion.vector_store.openai_client")
    def test_similarity_search(self, mock_openai, mock_qdrant_cls):
        from app.services.ingestion.vector_store import VectorStoreService

        mock_client = MagicMock()
        mock_client.get_collections.return_value = MagicMock(collections=[MagicMock(name="rag_chunks")])

        mock_result = MagicMock()
        mock_result.id = "p1"
        mock_result.score = 0.85
        mock_result.payload = {"chunk_text": "test", "chunk_id": "c1"}
        mock_client.search.return_value = [mock_result]
        mock_qdrant_cls.return_value = mock_client

        vs = VectorStoreService()
        results = vs.similarity_search([0.1] * 1536)
        assert len(results) == 1
        assert results[0]["score"] == 0.85

    @patch("app.services.ingestion.vector_store.QdrantClient")
    @patch("app.services.ingestion.vector_store.openai_client")
    def test_mmr_retrieval_empty(self, mock_openai, mock_qdrant_cls):
        from app.services.ingestion.vector_store import VectorStoreService

        mock_client = MagicMock()
        mock_client.get_collections.return_value = MagicMock(collections=[MagicMock(name="rag_chunks")])
        mock_client.search.return_value = []
        mock_qdrant_cls.return_value = mock_client

        vs = VectorStoreService()
        results = vs.mmr_retrieval([0.1] * 1536)
        assert results == []

    @patch("app.services.ingestion.vector_store.QdrantClient")
    @patch("app.services.ingestion.vector_store.openai_client")
    def test_mmr_retrieval_with_results(self, mock_openai, mock_qdrant_cls):
        from app.services.ingestion.vector_store import VectorStoreService

        mock_client = MagicMock()
        mock_client.get_collections.return_value = MagicMock(collections=[MagicMock(name="rag_chunks")])

        r1 = MagicMock()
        r1.id = "p1"
        r1.score = 0.9
        r1.payload = {"chunk_text": "first result", "chunk_id": "c1", "priority": 1.0}
        r1.vector = [0.1] * 1536

        r2 = MagicMock()
        r2.id = "p2"
        r2.score = 0.8
        r2.payload = {"chunk_text": "second result", "chunk_id": "c2", "priority": 1.0}
        r2.vector = [0.2] * 1536

        mock_client.search.return_value = [r1, r2]
        mock_qdrant_cls.return_value = mock_client

        vs = VectorStoreService()
        results = vs.mmr_retrieval([0.1] * 1536, top_k=2)
        assert len(results) >= 1

    @patch("app.services.ingestion.vector_store.QdrantClient")
    @patch("app.services.ingestion.vector_store.openai_client")
    def test_delete_by_document_id(self, mock_openai, mock_qdrant_cls):
        from app.services.ingestion.vector_store import VectorStoreService

        mock_client = MagicMock()
        mock_client.get_collections.return_value = MagicMock(collections=[MagicMock(name="rag_chunks")])
        mock_qdrant_cls.return_value = mock_client

        vs = VectorStoreService()
        vs.delete_by_document_id("doc-123")
        mock_client.delete.assert_called_once()

    def test_cosine_sim(self):
        from app.services.ingestion.vector_store import VectorStoreService

        a = np.array([1.0, 0.0, 0.0])
        b = np.array([1.0, 0.0, 0.0])
        assert VectorStoreService._cosine_sim(a, b) == pytest.approx(1.0)

        c = np.array([0.0, 1.0, 0.0])
        assert VectorStoreService._cosine_sim(a, c) == pytest.approx(0.0)

    def test_cosine_sim_zero_vector(self):
        from app.services.ingestion.vector_store import VectorStoreService

        a = np.array([0.0, 0.0])
        b = np.array([1.0, 0.0])
        assert VectorStoreService._cosine_sim(a, b) == 0.0

    def test_deduplicate(self):
        from app.services.ingestion.vector_store import VectorStoreService

        results = [
            {"payload": {"chunk_text": "hello world this is a test"}},
            {"payload": {"chunk_text": "hello world this is a test"}},  # exact duplicate
            {"payload": {"chunk_text": "completely different content here now"}},
        ]
        deduped = VectorStoreService._deduplicate(results)
        assert len(deduped) == 2

    def test_build_filter(self):
        from app.services.ingestion.vector_store import VectorStoreService

        f = VectorStoreService._build_filter({"document_id": "d1"})
        assert f is not None
        assert len(f.must) == 1

        f_none = VectorStoreService._build_filter(None)
        assert f_none is None

        f_empty = VectorStoreService._build_filter({})
        assert f_empty is None


# ─── S3 Service Tests ────────────────────────────────────────────────────────


class TestS3Service:
    @patch("app.services.ingestion.s3_service.boto3")
    def test_upload_document(self, mock_boto3):
        from app.services.ingestion.s3_service import S3Service
        import asyncio

        mock_client = MagicMock()
        mock_boto3.client.return_value = mock_client

        s3 = S3Service()

        mock_file = MagicMock()
        mock_file.read = AsyncMock(return_value=b"file content")
        mock_file.filename = "test.pdf"
        mock_file.content_type = "application/pdf"

        result = asyncio.get_event_loop().run_until_complete(
            s3.upload_document(mock_file, "user-123")
        )
        assert result["filename"] == "test.pdf"
        assert result["file_size_bytes"] == len(b"file content")
        mock_client.put_object.assert_called_once()

    @patch("app.services.ingestion.s3_service.boto3")
    def test_generate_presigned_url(self, mock_boto3):
        from app.services.ingestion.s3_service import S3Service

        mock_client = MagicMock()
        mock_client.generate_presigned_url.return_value = "https://s3.example.com/signed"
        mock_boto3.client.return_value = mock_client

        s3 = S3Service()
        url = s3.generate_presigned_url("documents/test.pdf")
        assert "signed" in url

    @patch("app.services.ingestion.s3_service.boto3")
    def test_download_to_bytes(self, mock_boto3):
        from app.services.ingestion.s3_service import S3Service

        mock_client = MagicMock()
        mock_body = MagicMock()
        mock_body.read.return_value = b"file bytes"
        mock_client.get_object.return_value = {"Body": mock_body}
        mock_boto3.client.return_value = mock_client

        s3 = S3Service()
        content = s3.download_to_bytes("documents/test.pdf")
        assert content == b"file bytes"

    @patch("app.services.ingestion.s3_service.boto3")
    def test_delete_object(self, mock_boto3):
        from app.services.ingestion.s3_service import S3Service

        mock_client = MagicMock()
        mock_boto3.client.return_value = mock_client

        s3 = S3Service()
        s3.delete_object("documents/test.pdf")
        mock_client.delete_object.assert_called_once()

    @patch("app.services.ingestion.s3_service.boto3")
    def test_delete_object_error_handled(self, mock_boto3):
        from app.services.ingestion.s3_service import S3Service
        from botocore.exceptions import ClientError

        mock_client = MagicMock()
        mock_client.delete_object.side_effect = ClientError(
            {"Error": {"Code": "NoSuchKey", "Message": "Not found"}}, "DeleteObject"
        )
        mock_boto3.client.return_value = mock_client

        s3 = S3Service()
        # Should not raise
        s3.delete_object("nonexistent-key")

    @patch("app.services.ingestion.s3_service.boto3")
    def test_generate_key_uniqueness(self, mock_boto3):
        from app.services.ingestion.s3_service import S3Service

        mock_boto3.client.return_value = MagicMock()
        s3 = S3Service()
        k1 = s3._generate_key("user1", "test.pdf")
        k2 = s3._generate_key("user1", "test.pdf")
        assert k1 != k2  # UUID makes it unique
        assert k1.endswith(".pdf")
        assert "user1" in k1


# ─── Session Service Tests ───────────────────────────────────────────────────


class TestSessionService:
    @pytest.mark.asyncio
    async def test_create_conversation(self, db_session: AsyncSession):
        from app.services.session.session_service import SessionService

        # Create a test user
        user = User(email=f"sess_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash", full_name="Test")
        db_session.add(user)
        await db_session.flush()

        svc = SessionService(db_session)
        conv = await svc.create_conversation(user, title="Test Conv")
        assert conv.title == "Test Conv"
        assert conv.user_id == user.id
        assert conv.session_type == SessionType.PERMANENT

    @pytest.mark.asyncio
    async def test_create_temporary_conversation(self, db_session: AsyncSession):
        from app.services.session.session_service import SessionService

        user = User(email=f"temp_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash")
        db_session.add(user)
        await db_session.flush()

        svc = SessionService(db_session)
        conv = await svc.create_conversation(user, session_type=SessionType.TEMPORARY)
        assert conv.session_type == SessionType.TEMPORARY
        assert conv.expires_at is not None

    @pytest.mark.asyncio
    async def test_get_conversation(self, db_session: AsyncSession):
        from app.services.session.session_service import SessionService

        user = User(email=f"get_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash")
        db_session.add(user)
        await db_session.flush()

        svc = SessionService(db_session)
        conv = await svc.create_conversation(user, title="Findable")
        await db_session.commit()

        found = await svc.get_conversation(conv.id, user)
        assert found is not None
        assert found.title == "Findable"

    @pytest.mark.asyncio
    async def test_get_conversation_wrong_user(self, db_session: AsyncSession):
        from app.services.session.session_service import SessionService

        user1 = User(email=f"u1_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash")
        user2 = User(email=f"u2_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash")
        db_session.add_all([user1, user2])
        await db_session.flush()

        svc = SessionService(db_session)
        conv = await svc.create_conversation(user1, title="Private")
        await db_session.commit()

        found = await svc.get_conversation(conv.id, user2)
        assert found is None

    @pytest.mark.asyncio
    async def test_add_message(self, db_session: AsyncSession):
        from app.services.session.session_service import SessionService

        user = User(email=f"msg_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash")
        db_session.add(user)
        await db_session.flush()

        svc = SessionService(db_session)
        conv = await svc.create_conversation(user)
        msg = await svc.add_message(conv.id, MessageRole.USER, "Hello!")
        assert msg.content == "Hello!"
        assert msg.role == MessageRole.USER

    @pytest.mark.asyncio
    async def test_get_recent_messages(self, db_session: AsyncSession):
        from app.services.session.session_service import SessionService

        user = User(email=f"recent_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash")
        db_session.add(user)
        await db_session.flush()

        svc = SessionService(db_session)
        conv = await svc.create_conversation(user)
        for i in range(5):
            await svc.add_message(conv.id, MessageRole.USER, f"Message {i}")
        await db_session.commit()

        recent = await svc.get_recent_messages(conv.id, limit=3)
        assert len(recent) == 3
        assert recent[0]["role"] == "user"

    @pytest.mark.asyncio
    async def test_list_conversations(self, db_session: AsyncSession):
        from app.services.session.session_service import SessionService

        user = User(email=f"list_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash")
        db_session.add(user)
        await db_session.flush()

        svc = SessionService(db_session)
        await svc.create_conversation(user, title="Conv 1")
        await svc.create_conversation(user, title="Conv 2")
        await db_session.commit()

        convs = await svc.list_conversations(user)
        assert len(convs) >= 2

    @pytest.mark.asyncio
    async def test_get_or_update_summary_no_trigger(self, db_session: AsyncSession):
        from app.services.session.session_service import SessionService

        user = User(email=f"sum_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash")
        db_session.add(user)
        await db_session.flush()

        svc = SessionService(db_session)
        conv = await svc.create_conversation(user)
        await svc.add_message(conv.id, MessageRole.USER, "Hi")
        await db_session.commit()

        summary = await svc.get_or_update_summary(conv)
        assert summary is None  # Not enough messages to trigger

    @pytest.mark.asyncio
    async def test_expire_stale_sessions(self, db_session: AsyncSession):
        from app.services.session.session_service import SessionService

        user = User(email=f"expire_{uuid.uuid4().hex[:6]}@test.com", hashed_password="hash")
        db_session.add(user)
        await db_session.flush()

        # Create an expired temp conversation
        conv = Conversation(
            user_id=user.id,
            title="Expired",
            session_type=SessionType.TEMPORARY,
            is_active=True,
            expires_at=datetime.now(timezone.utc) - timedelta(hours=1),
        )
        db_session.add(conv)
        await db_session.commit()

        svc = SessionService(db_session)
        count = await svc.expire_stale_sessions()
        assert count >= 1


# ─── Auth Service Integration Tests ──────────────────────────────────────────


class TestAuthServiceDB:
    @pytest.mark.asyncio
    async def test_register_and_authenticate(self, db_session: AsyncSession):
        from app.services.auth.auth_service import AuthService

        svc = AuthService(db_session)
        user = await svc.register_user(f"auth_{uuid.uuid4().hex[:6]}@test.com", "password123")
        assert user.role == UserRole.STUDENT

        authed = await svc.authenticate(user.email, "password123")
        assert authed.id == user.id

    @pytest.mark.asyncio
    async def test_authenticate_wrong_password(self, db_session: AsyncSession):
        from app.services.auth.auth_service import AuthService
        from fastapi import HTTPException

        svc = AuthService(db_session)
        email = f"wrong_{uuid.uuid4().hex[:6]}@test.com"
        await svc.register_user(email, "correct")
        await db_session.commit()

        with pytest.raises(HTTPException) as exc_info:
            await svc.authenticate(email, "wrong")
        assert exc_info.value.status_code == 401

    @pytest.mark.asyncio
    async def test_register_duplicate(self, db_session: AsyncSession):
        from app.services.auth.auth_service import AuthService
        from fastapi import HTTPException

        svc = AuthService(db_session)
        email = f"dup_{uuid.uuid4().hex[:6]}@test.com"
        await svc.register_user(email, "pass12345")
        await db_session.commit()

        with pytest.raises(HTTPException) as exc_info:
            await svc.register_user(email, "pass12345")
        assert exc_info.value.status_code == 409

    @pytest.mark.asyncio
    async def test_store_and_rotate_refresh_token(self, db_session: AsyncSession):
        from app.services.auth.auth_service import AuthService, create_refresh_token

        svc = AuthService(db_session)
        email = f"rotate_{uuid.uuid4().hex[:6]}@test.com"
        user = await svc.register_user(email, "pass12345")
        refresh = create_refresh_token(str(user.id))
        await svc.store_refresh_token(user.id, refresh)
        await db_session.commit()

        result_user, new_access, new_refresh = await svc.rotate_refresh_token(refresh)
        assert result_user.id == user.id
        assert new_access != refresh
        assert new_refresh != refresh

    @pytest.mark.asyncio
    async def test_blocked_user_cannot_authenticate(self, db_session: AsyncSession):
        from app.services.auth.auth_service import AuthService
        from fastapi import HTTPException

        svc = AuthService(db_session)
        email = f"blocked_{uuid.uuid4().hex[:6]}@test.com"
        user = await svc.register_user(email, "pass12345")
        user.is_active = False
        await db_session.commit()

        with pytest.raises(HTTPException) as exc_info:
            await svc.authenticate(email, "pass12345")
        assert exc_info.value.status_code == 403


# ─── Document Extractor Tests ────────────────────────────────────────────────


class TestExtractorAdvanced:
    def test_plain_text_extractor(self):
        from app.services.ingestion.extractor import PlainTextExtractor

        ext = PlainTextExtractor()
        result = ext.extract(b"Simple text content")
        assert "Simple text" in result.text
        assert result.is_ocr is False

    def test_docx_extractor(self):
        from app.services.ingestion.extractor import DocxExtractor
        import io
        from docx import Document as DocxDocument

        # Create a real DOCX in memory
        doc = DocxDocument()
        doc.add_paragraph("Test paragraph one")
        doc.add_paragraph("Test paragraph two")
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        ext = DocxExtractor()
        result = ext.extract(content)
        assert "Test paragraph one" in result.text
        assert "Test paragraph two" in result.text
        assert result.is_ocr is False

    def test_document_extractor_factory_txt(self):
        from app.services.ingestion.extractor import DocumentExtractor

        ext = DocumentExtractor()
        result = ext.extract(b"Hello world", "readme.txt")
        assert "Hello world" in result.text

    def test_document_extractor_factory_md(self):
        from app.services.ingestion.extractor import DocumentExtractor

        ext = DocumentExtractor()
        result = ext.extract(b"# Title\nContent here", "notes.md")
        assert "Title" in result.text

    def test_document_extractor_factory_docx(self):
        from app.services.ingestion.extractor import DocumentExtractor
        import io
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Test content")
        buf = io.BytesIO()
        doc.save(buf)

        ext = DocumentExtractor()
        result = ext.extract(buf.getvalue(), "test.docx")
        assert "Test content" in result.text


# ─── Celery Task Tests ───────────────────────────────────────────────────────


class TestCeleryTasks:
    @patch("app.services.ingestion.vector_store.VectorStoreService")
    @patch("app.services.ingestion.s3_service.S3Service")
    @patch("app.services.ingestion.extractor.DocumentExtractor")
    @patch("app.services.ingestion.chunker.ChunkerFactory")
    def test_process_document_success(self, mock_chunker_factory,
                                      mock_extractor_cls, mock_s3_cls, mock_vs_cls):
        from app.tasks.celery_app import process_document, SyncSession
        from app.services.ingestion.chunker import Chunk

        # Mock DB session
        mock_db = MagicMock()

        # Mock document record
        mock_doc = MagicMock()
        mock_doc.id = uuid.uuid4()
        mock_doc.s3_key = "documents/test.pdf"
        mock_doc.filename = "test.pdf"
        mock_doc.document_type = DocumentType.LEARNING_MATERIAL
        mock_doc.subject_name = None
        mock_doc.subject_code = None
        mock_db.execute.return_value.scalar_one_or_none.return_value = mock_doc

        # Mock S3
        mock_s3 = MagicMock()
        mock_s3.download_to_bytes.return_value = b"file content"
        mock_s3_cls.return_value = mock_s3

        # Mock extractor
        mock_extraction = MagicMock()
        mock_extraction.text = "Extracted text content"
        mock_extraction.page_count = 1
        mock_extraction.is_ocr = False
        mock_extractor_cls.return_value.extract.return_value = mock_extraction

        # Mock chunker
        mock_chunker = MagicMock()
        mock_chunks = [
            Chunk(chunk_text="Chunk 1", chunk_index=0, document_type=DocumentType.LEARNING_MATERIAL),
            Chunk(chunk_text="Chunk 2", chunk_index=1, document_type=DocumentType.LEARNING_MATERIAL),
        ]
        mock_chunker.chunk.return_value = mock_chunks
        mock_chunker_factory.get_chunker.return_value = (mock_chunker, {})

        # Mock vector store
        mock_vs = MagicMock()
        mock_vs.upsert_chunks.return_value = ["pid1", "pid2"]
        mock_vs_cls.return_value = mock_vs

        # Run task (bypass Celery, mock the DB session)
        with patch("app.tasks.celery_app.SyncSession", return_value=mock_db):
            process_document(str(mock_doc.id))

        mock_db.commit.assert_called()
        assert mock_doc.status == DocumentStatus.INDEXED

    def test_process_document_not_found(self):
        from app.tasks.celery_app import process_document

        mock_db = MagicMock()
        mock_db.execute.return_value.scalar_one_or_none.return_value = None

        with patch("app.tasks.celery_app.SyncSession", return_value=mock_db):
            process_document("nonexistent-id")
        # No commit should be called since doc not found
        mock_db.commit.assert_not_called()


# ─── Cache Service Async Tests ───────────────────────────────────────────────


class TestCacheServiceAsync:
    @pytest.mark.asyncio
    @patch("app.utils.cache.get_redis")
    async def test_get_cached_response_miss(self, mock_get_redis):
        from app.utils.cache import CacheService

        mock_redis = MagicMock()
        mock_redis.get = AsyncMock(return_value=None)
        mock_get_redis.return_value = mock_redis

        cache = CacheService()
        result = await cache.get_cached_response("query", "conv-1")
        assert result is None

    @pytest.mark.asyncio
    @patch("app.utils.cache.get_redis")
    async def test_get_cached_response_hit(self, mock_get_redis):
        from app.utils.cache import CacheService

        cached_data = {"answer": "cached answer", "sources": []}
        mock_redis = MagicMock()
        mock_redis.get = AsyncMock(return_value=json.dumps(cached_data))
        mock_get_redis.return_value = mock_redis

        cache = CacheService()
        result = await cache.get_cached_response("query", "conv-1")
        assert result["answer"] == "cached answer"

    @pytest.mark.asyncio
    @patch("app.utils.cache.get_redis")
    async def test_set_cached_response(self, mock_get_redis):
        from app.utils.cache import CacheService

        mock_redis = MagicMock()
        mock_redis.setex = AsyncMock()
        mock_get_redis.return_value = mock_redis

        cache = CacheService()
        await cache.set_cached_response("query", "conv-1", {"answer": "test"})
        mock_redis.setex.assert_called_once()

    @pytest.mark.asyncio
    @patch("app.utils.cache.get_redis")
    async def test_get_set_delete(self, mock_get_redis):
        from app.utils.cache import CacheService

        mock_redis = MagicMock()
        mock_redis.get = AsyncMock(return_value="value")
        mock_redis.setex = AsyncMock()
        mock_redis.delete = AsyncMock()
        mock_get_redis.return_value = mock_redis

        cache = CacheService()
        await cache.set("key", "value")
        result = await cache.get("key")
        assert result == "value"
        await cache.delete("key")
        mock_redis.delete.assert_called_once()

    @pytest.mark.asyncio
    @patch("app.utils.cache.get_redis")
    async def test_increment(self, mock_get_redis):
        from app.utils.cache import CacheService

        mock_redis = MagicMock()
        mock_pipe = MagicMock()
        mock_pipe.incr = AsyncMock()
        mock_pipe.expire = AsyncMock()
        mock_pipe.execute = AsyncMock(return_value=[5])
        mock_redis.pipeline.return_value = mock_pipe
        mock_get_redis.return_value = mock_redis

        cache = CacheService()
        count = await cache.increment("rate:user1")
        assert count == 5


# ─── Logging Configuration Tests ─────────────────────────────────────────────


class TestLogging:
    def test_configure_logging(self):
        from app.core.logging import configure_logging
        # Should not raise
        configure_logging()


# ─── Rate Limiter Tests ──────────────────────────────────────────────────────


class TestRateLimiterDetail:
    def test_get_identifier_with_user(self):
        from app.core.rate_limiter import _get_identifier

        request = MagicMock()
        request.state.user = MagicMock()
        request.state.user.id = "user-123"
        result = _get_identifier(request)
        assert result == "user:user-123"

    def test_get_identifier_without_user(self):
        from app.core.rate_limiter import _get_identifier

        request = MagicMock()
        request.state = MagicMock(spec=[])  # No 'user' attribute
        request.client.host = "127.0.0.1"
        result = _get_identifier(request)
        assert result == "127.0.0.1"
